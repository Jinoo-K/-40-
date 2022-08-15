# pip install openpyxl
# pip install python-docx
# pip install docx2pdf

import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r"12. 엑셀의 정보를 불러와 수료증 자동 생성\수료증양식.docx")

style = doc.styles["Normal"]
style.font.name = "나눔고딕"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run(" 제 2022-0001 호\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(15)

para = doc.add_paragraph()
run = para.add_run("\n")
run = para.add_run("수 료 증")
run.font.name = "나눔고딕"
run.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run("\n")
run = para.add_run(" 성 명 : 김진우\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(15)
run = para.add_run(" 생 년 월 일 : 1991.01.23\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(15)
run = para.add_run(" 교 육 과 정 : 파이썬과 40개의 작품들\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(15)
run = para.add_run (" 교 육 날 짜 : 2022.08.08 ~ 2022.08.31\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(15)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run(" 위 사람은 파이썬과 40개의 작품들 교육 과정을\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)
run = para.add_run("이수하였으므로 이 증서를 수여합니다.\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run("2022.08.31")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run("파이썬 교육 기관장")
run.font.name = "나눔고딕"
run.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r"12. 엑셀의 정보를 불러와 수료증 자동 생성\수료증결과.docx")